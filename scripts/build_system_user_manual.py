from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public" / "downloads" / "DomainWatch-User-Manual-v2.3.docx"
ASSETS = ROOT / "public" / "help"
FONT = "Kanit"
NAVY = "172554"
BLUE = "2458E6"
INK = "1E293B"
MUTED = "64748B"
LIGHT = "F5F7FC"
GREEN = "059669"
AMBER = "D97706"
RED = "DC2626"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_image_alt(run, alt_text):
    drawing = run._element.find(qn("w:drawing"))
    if drawing is None:
        return
    for doc_pr in drawing.iter(qn("wp:docPr")):
        doc_pr.set("descr", alt_text)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("DomainWatch • ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    heading_specs = {
        "Title": (29, NAVY, 0, 10),
        "Subtitle": (13, MUTED, 0, 8),
        "Heading 1": (18, NAVY, 16, 8),
        "Heading 2": (14, BLUE, 12, 6),
        "Heading 3": (11.5, INK, 9, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.text = "DOMAINWATCH  /  USER GUIDE"
    header.runs[0].font.name = FONT
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    add_page_number(section.footer.paragraphs[0])


def force_fonts(doc):
    def walk(parent):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from walk(cell)

    parts = [doc]
    for section in doc.sections:
        parts.extend([section.header, section.footer])
    for part in parts:
        for paragraph in walk(part):
            for run in paragraph.runs:
                run.font.name = FONT
                r_pr = run._element.get_or_add_rPr()
                r_fonts = r_pr.rFonts
                if r_fonts is None:
                    r_fonts = OxmlElement("w:rFonts")
                    r_pr.insert(0, r_fonts)
                for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                    r_fonts.set(qn(f"w:{key}"), FONT)


def add_title(doc, title, subtitle=None):
    paragraph = doc.add_heading(title, level=1)
    paragraph.paragraph_format.page_break_before = True
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
        p.runs[0].font.size = Pt(10)


def add_bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        doc.add_paragraph(item, style=style)


def add_callout(doc, title, text, color="EFF6FF", title_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.65)
    set_cell_fill(cell, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(title_color)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_picture(doc, filename, caption, width=6.45):
    path = ASSETS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    set_image_alt(run, caption)
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    for r in cp.runs:
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(MUTED)
        r.italic = True


def add_status_table(doc):
    rows = [
        ("สีเขียว — ใช้งานได้", "ตอบกลับและผ่านเกณฑ์ ระบบถือว่าใช้งานได้", GREEN),
        ("สีเหลือง — โหลดช้า", "เปิดได้ แต่ใช้เวลานาน ยังไม่ถือว่าเว็บล่ม", AMBER),
        ("สีแดง — ใช้ไม่ได้", "ยืนยันผิดปกติครบตามเกณฑ์ ต้องตรวจสอบหรือแก้ลิงก์", RED),
        ("เทา — ยังไม่ทราบ", "ยังไม่มีผลตรวจล่าสุด เครื่องออฟไลน์ หรือข้อมูลยังไม่ครบ", MUTED),
        ("พักชั่วคราว", "หยุดตรวจและแจ้งเตือน แต่เก็บประวัติเดิมไว้", BLUE),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "สถานะ"
    headers[1].text = "ความหมายและสิ่งที่ควรทำ"
    set_repeat_table_header(table.rows[0])
    for cell, width in zip(headers, (2.05, 4.6)):
        set_cell_width(cell, width)
        set_cell_fill(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    for label, meaning, color in rows:
        cells = table.add_row().cells
        set_cell_width(cells[0], 2.05)
        set_cell_width(cells[1], 4.6)
        cells[0].text = label
        cells[1].text = meaning
        cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(color)
        cells[0].paragraphs[0].runs[0].bold = True


def add_role_table(doc):
    data = [
        ("ADMIN", "ทุกหน้าและจัดการผู้ใช้", "ทุกบริษัท"),
        ("ADMIN_LEAD", "จัดการเคส/ลิงก์ และดู KPI", "ทุกบริษัท"),
        ("ADMIN_COMPANY", "จัดการเฉพาะบริษัทที่มอบหมาย ไม่เห็น KPI", "บริษัทที่ผูกไว้"),
        ("IT", "รับเคสและใส่ลิงก์สำรอง", "ทุกบริษัท"),
        ("MANAGEMENT", "อ่าน Dashboard รายงาน และ KPI", "ทุกบริษัท"),
        ("SITE_STAFF", "ดูหน้าเครื่องตรวจ ดาวน์โหลด APK และผูกเครื่อง", "ตามสิทธิ์ที่กำหนด"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (1.45, 3.65, 1.55)
    for index, (heading, width) in enumerate(zip(("บทบาท", "ทำอะไรได้", "ขอบเขต"), widths)):
        cell = table.rows[0].cells[index]
        cell.text = heading
        set_cell_width(cell, width)
        set_cell_fill(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    for role, permission, scope in data:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, (role, permission, scope), widths):
            cell.text = value
            set_cell_width(cell, width)
        cells[0].paragraphs[0].runs[0].bold = True


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(64)
    r = p.add_run("D")
    r.bold = True
    r.font.size = Pt(52)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    title = doc.add_paragraph("คู่มือการใช้งาน DomainWatch", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("สำหรับผู้ดูแลระบบ แอดมิน ไอที ผู้บริหาร และพนักงานหน้าไซต์", style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture(doc, "login-qr.png", "สแกน QR เพื่อเปิดหน้าเข้าสู่ระบบ DomainWatch", width=2.25)
    p = doc.add_paragraph("https://domain-watch-app-sandy.vercel.app/login")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    p.runs[0].bold = True
    p2 = doc.add_paragraph("ฉบับ 2.3 • 1 กันยายน 2569 • เอกสารสำหรับส่งมอบงาน")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(9.5)
    p2.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_callout(doc, "ข้อควรจำ", "ใช้บัญชีของตนเองเท่านั้น ห้ามส่งรหัสผ่านหรือ QR ผูกเครื่องลงในกลุ่มสาธารณะ", "FEF3C7", "92400E")

    add_title(doc, "สารบัญและเส้นทางใช้งาน", "ใช้หัวข้อต่อไปนี้เป็นลำดับสอนงานพนักงานใหม่")
    add_bullets(doc, [
        "1. เข้าใช้ครั้งแรก — เปิดลิงก์หรือสแกน QR และเข้าสู่ระบบ",
        "2. เมนูและสิทธิ์ — เห็นเฉพาะหน้าที่บทบาทของตนอนุญาต",
        "3. แดชบอร์ด — อ่านยอดสถานะและรายการที่ต้องดำเนินการ",
        "4. Master Data — เพิ่ม แก้ไข พักเฝ้าดู และจัดการลิงก์สำรอง",
        "5. เหตุการณ์ — แยกปัญหาระบบกลางกับปัญหาเครือข่ายซิม",
        "6. เครื่องตรวจเครือข่าย — ติดตั้ง APK เปิด VPN กรุงเทพฯ สร้าง QR และดูผลรายเครื่อง",
        "7. รายงานรอบวัน / KPI / Telegram — ตรวจสอบและส่งมอบรายงาน",
        "8. วิธีแก้ปัญหาที่พบบ่อยและรายการตรวจรับงาน",
    ])
    add_callout(doc, "Flow มาตรฐาน", "รับแจ้ง → เปิดเคส → ตรวจสาเหตุ → แก้ลิงก์หรือใส่ลิงก์สำรอง → รอตรวจยืนยัน → ปิดเคส → ตรวจในรายงานรอบวัน", "EFF6FF", BLUE)

    add_title(doc, "1. เข้าใช้ระบบครั้งแรก")
    doc.add_heading("วิธี A — เปิดจากลิงก์", level=2)
    add_bullets(doc, [
        "เปิด Chrome หรือ Safari แล้วเข้า https://domain-watch-app-sandy.vercel.app/login",
        "กรอกชื่อผู้ใช้และรหัสผ่านที่แอดมินมอบให้ แล้วกด เข้าสู่ระบบ",
        "หากระบบขอเปลี่ยนรหัสผ่าน ให้ตั้งรหัสใหม่ที่ไม่ซ้ำกับบัญชีอื่น",
        "หลังเข้าได้ ให้ตรวจชื่อและบทบาทมุมขวาบนก่อนเริ่มงาน",
    ], numbered=True)
    doc.add_heading("วิธี B — สแกน QR", level=2)
    add_picture(doc, "login-qr.png", "QR สำหรับเปิดหน้าเข้าสู่ระบบ — สแกนด้วยกล้องโทรศัพท์", width=2.35)
    add_bullets(doc, [
        "เปิดกล้องโทรศัพท์หรือเมนูสแกน QR แล้วเล็งให้เห็นกรอบทั้งหมด",
        "แตะข้อความแจ้งเตือนเพื่อเปิดหน้าเข้าสู่ระบบในเบราว์เซอร์",
        "ตรวจโดเมนให้เป็น domain-watch-app-sandy.vercel.app ก่อนกรอกรหัสผ่าน",
        "QR นี้เปิดได้เพียงหน้าเข้าสู่ระบบ ไม่ได้บรรจุชื่อผู้ใช้หรือรหัสผ่าน",
    ], numbered=True)
    add_callout(doc, "เข้าไม่ได้ทำอย่างไร", "ตรวจอินเทอร์เน็ตและตัวสะกดก่อน หากลืมรหัสผ่านให้ติดต่อ ADMIN ผู้ใช้ที่ยังไม่เข้าสู่ระบบจะไม่เห็นเมนูภายใน", "FEE2E2", RED)

    add_title(doc, "2. เมนูและสิทธิ์ผู้ใช้")
    add_role_table(doc)
    doc.add_paragraph("หลักการ: สิทธิ์แค่ไหน เห็นและทำได้แค่นั้น ทั้งเมนู ปุ่ม และข้อมูลที่ API ส่งกลับ")
    doc.add_heading("เปิดเมนูบนโทรศัพท์", level=2)
    add_picture(doc, "mobile-navigation.png", "หน้า Dashboard บนมือถือ — แตะปุ่มเมนูมุมซ้ายบน", width=3.0)
    add_picture(doc, "mobile-menu-open.png", "แถบเมนูบนมือถือ — เลือกหน้าที่ต้องการแล้วเมนูจะปิดอัตโนมัติ", width=3.0)

    add_title(doc, "3. แดชบอร์ดและการอ่านยอด")
    add_picture(doc, "dashboard.jpg", "Dashboard แสดงภาพรวมของลิงก์ที่เฝ้าดูและงานที่ต้องดำเนินการ", width=5.65)
    add_status_table(doc)
    doc.add_heading("ยอดที่ต้องอ่านให้ถูก", level=2)
    add_bullets(doc, [
        "ลิงก์ทั้งหมด คือจำนวนระเบียนใน Master Data รวมทั้งลิงก์ LINE ที่ไม่ตรวจหน้าเว็บ",
        "ใช้งานได้ / โหลดช้า / ใช้ไม่ได้ นับจากลิงก์หน้าเว็บที่เฝ้าดูและมีผลตรวจล่าสุด",
        "URL จริง คือจำนวน URL ไม่ซ้ำ ส่วนเคสเปิดนับตามบริษัทและห้อง LINE จึงอาจมากกว่า URL จริง",
        "ลิงก์สำรองครบ แสดงจำนวนระเบียนที่เตรียมลิงก์สำรองไว้ ไม่ใช่จำนวน URL ไม่ซ้ำ",
        "กดการ์ดสีหรือ ดูรายละเอียด เพื่อเปิดรายการที่ใช้สร้างยอดนั้นทันที",
    ])

    add_title(doc, "4. บริษัท ห้อง LINE และ Master Data ลิงก์")
    add_picture(doc, "companies.jpg", "หน้าบริษัท / ห้อง LINE — ใช้ค้นหาและเปิดรายละเอียดแต่ละห้อง", width=6.35)
    add_picture(doc, "links.jpg", "Master Data ลิงก์ — กรองตามบริษัท ห้อง หมวด สถานะ และแหล่งตรวจ", width=6.35)
    doc.add_heading("เพิ่มหรือแก้ไขลิงก์", level=2)
    add_bullets(doc, [
        "เลือกบริษัทและห้อง LINE ให้ถูก เพราะลิงก์เดียวกันในคนละห้องถือเป็นคนละรายการและต้องแจ้งแยกกัน",
        "กรอกชื่อลิงก์ URL หมวดหมู่ และลิงก์สำรองถ้ามี",
        "เปิด เฝ้าดู เมื่อเป็นหน้าเว็บที่ต้องตรวจ ปิดเมื่อเป็นลิงก์ LINE ที่ตรวจไม่ได้",
        "ถ้าเว็บหยุดดำเนินการชั่วคราว ให้เลือก พักชั่วคราว ประวัติและเคสเดิมจะยังอยู่ แต่หยุดตรวจและแจ้งเตือน",
        "บันทึกแล้วตรวจวันที่เช็คล่าสุดและสถานะในตาราง",
    ], numbered=True)
    add_callout(doc, "กรณี Redirect ของเครือข่าย", "เก็บ URL หลักไว้เป็นหลักฐาน ใส่ปลายทางที่เครือข่ายใช้งานได้ในช่องลิงก์สำรอง เมื่อเครื่องซิมยืนยันว่าลิงก์สำรองใช้งานได้ ระบบถือว่าบริการกลับมาใช้ได้และปิดเคสเครือข่ายเดิมได้", "ECFDF5", GREEN)
    doc.add_heading("ตรวจ LINE OA และดึงลิงก์จาก Rich Menu", level=2)
    add_bullets(doc, [
        "เข้า LINE Developers Console เลือก Provider และ Messaging API channel ของ OA จากแท็บ Messaging API ให้ออก Channel access token แล้วเก็บเป็นความลับ",
        "ใน DomainWatch ไปที่ บริษัท / ห้อง LINE กางบริษัท แล้วกด ⚙️ OA ที่ห้องเป้าหมาย",
        "วาง Channel Access Token กรอกชื่อ OA ที่คาดหวัง และกด บันทึก ก่อน",
        "กด ตรวจ OA ตอนนี้ เพื่อตรวจว่า token ใช้ได้ ชื่อบัญชีตรง และมีรูปโปรไฟล์หรือไม่",
        "กด ดึงลิงก์จาก Rich Menu ระบบจะเพิ่ม URL ใหม่ในห้องนั้นเป็นหมวด ริชเมนู และข้าม URL ที่มีอยู่แล้ว",
        "ไปที่ Master Data กรองบริษัทและห้อง ตรวจชื่อ URL สถานะเฝ้าดู และเติมลิงก์สำรองก่อนใช้งานจริง",
    ], numbered=True)
    add_callout(doc, "ข้อจำกัดของ LINE", "Messaging API ดึงรายการ Rich Menu ได้เฉพาะเมนูที่สร้างผ่าน Messaging API เท่านั้น เมนูที่สร้างใน LINE Official Account Manager อาจดึงไม่ได้ หากระบบแจ้งว่าไม่พบลิงก์ ให้ตรวจแหล่งที่สร้างเมนูและเพิ่ม URL ใน Master Data ด้วยตนเอง", "FEF3C7", "92400E")
    add_callout(doc, "รักษา Token ให้ปลอดภัย", "Channel Access Token มีสิทธิ์เรียก Messaging API ห้ามส่งในกลุ่มหรือแนบในภาพหน้าจอ หากสงสัยว่ารั่วให้ revoke หรือออก token ใหม่ใน LINE Developers Console", "FEE2E2", RED)

    add_title(doc, "5. เหตุการณ์และสถานะการจัดการ")
    doc.add_heading("แหล่งตรวจ 2 แบบ", level=2)
    add_bullets(doc, [
        "ระบบกลาง: ตรวจจากเซิร์ฟเวอร์ ใช้ดูการตอบสนองโดยรวม",
        "เครือข่ายซิม: ตรวจผ่านโทรศัพท์และซิมจริง ใช้ยืนยันปัญหาเฉพาะเครือข่าย",
    ])
    doc.add_heading("สถานะเคส", level=2)
    add_bullets(doc, [
        "เปิด (รอจัดการ): ยืนยันปัญหาแล้วและยังไม่ได้แก้",
        "ปรับแก้แล้ว · รอตรวจยืนยัน: ผู้ใช้แก้ URL หรือเพิ่มลิงก์สำรองแล้ว ระบบกำลังรอผลรอบใหม่",
        "จัดการแล้ว: ระบบยืนยันผลปกติหรือผู้มีสิทธิ์ปิดเคสพร้อมบันทึกเหตุผล",
        "พักชั่วคราว: หยุดตรวจรายการนั้น แต่ประวัติเดิมไม่ถูกลบ",
    ])
    doc.add_heading("วิธีแก้จากหน้าการ์ดเหตุการณ์", level=2)
    add_bullets(doc, [
        "กด แก้ลิงก์ตรงนี้ ไม่ต้องย้อนกลับไป Master Data",
        "แก้ URL หลักหรือใส่ลิงก์สำรอง แล้วกดบันทึก",
        "ตรวจให้การ์ดเปลี่ยนเป็น ปรับแก้แล้ว · รอตรวจยืนยัน",
        "เมื่อผลผ่านตามเกณฑ์ เคสจะไปที่ประวัติทั้งหมดและแสดงเวลาที่แก้เสร็จ",
    ], numbered=True)

    add_title(doc, "6. เครื่องตรวจเครือข่ายมือถือ")
    add_picture(doc, "mobile-agents.png", "หน้าเครื่องตรวจเครือข่ายบนมือถือ — ดาวน์โหลด APK สร้าง QR และเลือกดูเครื่อง", width=3.05)
    doc.add_heading("ติดตั้งและผูกเครื่อง", level=2)
    add_bullets(doc, [
        "ใส่ซิม เปิด Mobile data และปิด Wi‑Fi ระหว่างตรวจรับครั้งแรก",
        "แอดมินสร้างชื่อเครื่องและ QR ติดตั้งครบชุด ซึ่งใช้ครั้งเดียวภายใน 30 นาที",
        "สแกน QR จากมือถือ หน้าเดียวมีทั้งดาวน์โหลด APK 1.0.6 และปุ่มเปิดแอปผูกเครื่อง",
        "ติดตั้งเสร็จแล้วกลับหน้าเดิม กด ติดตั้งหรือเปิดแอป แล้วผูกเครื่อง ระบบกรอก URL, token, ชื่อเครื่อง, ค่าย และโหมดตรวจให้อัตโนมัติ",
        "กด อนุญาตให้ทำงานเบื้องหลัง และยกเว้นการประหยัดแบตเตอรี่",
        "กด เริ่มตรวจตลอดเวลา แล้วรอให้หน้าเว็บแสดง ออนไลน์ และมีผลรอบแรก",
    ], numbered=True)
    doc.add_heading("แอป VPN ที่ใช้", level=2)
    add_bullets(doc, [
        "DomainWatch Agent เป็นแอปตรวจ URL ของระบบ ไม่ใช่แอป VPN",
        "แนะนำ Surfshark VPN สำหรับขั้นตอนนี้ เพราะมีแอป Android และเลือก Thailand → Bangkok ได้",
        "ดาวน์โหลดจาก Google Play: https://play.google.com/store/apps/details?id=com.surfshark.vpnclient.android",
        "Surfshark เป็นบริการสมัครสมาชิก ต้องใช้บัญชี/แพ็กเกจของบริษัท DomainWatch ไม่ได้รวมค่าบริการ VPN",
    ])
    add_picture(doc, "vpn-mobile-flow.png", "Flow เปิด Surfshark VPN กรุงเทพก่อนเริ่ม DomainWatch Agent", width=6.55)
    doc.add_heading("วิธีเปิด VPN กรุงเทพฯ", level=2)
    add_bullets(doc, [
        "ปิด Wi‑Fi และเปิด Mobile data ของซิม TRUE",
        "เปิด Surfshark และเข้าสู่ระบบด้วยบัญชีของบริษัท",
        "ค้นหา Thailand แล้วเลือก Bangkok จากนั้นกด Connect",
        "อนุญาตคำขอสร้าง VPN ของ Android ครั้งแรก และรอจนขึ้น Connected/มีรูปกุญแจ",
        "เปิด DomainWatch Agent กด เริ่มตรวจตลอดเวลา และรอประมาณ 5 นาที",
        "ในหน้าเครื่องตรวจเลือก VPN ที่เปิดอยู่บนมือถือ แล้วตรวจเมือง/ภูมิภาค/ประเทศของ IP ทางออก",
    ], numbered=True)
    add_callout(doc, "ตำแหน่งที่ระบบแสดง", "เป็นตำแหน่งโดยประมาณของ Public IP เซิร์ฟเวอร์ VPN ไม่ใช่ GPS หรือที่ตั้งจริงของมือถือ ฐานข้อมูล IP อาจแสดงเพียง Bangkok/Thailand ไม่ถึงระดับสยาม", "EFF6FF", BLUE)
    add_callout(doc, "เมื่อ VPN หลุด", "แอปจะไม่ใช้ผลจากซิมมาแทนในโหมด VPN ให้เชื่อม Surfshark ใหม่และรอรอบถัดไป ทั้ง Surfshark และ DomainWatch ต้องได้รับอนุญาตทำงานเบื้องหลัง", "FEF3C7", "92400E")
    add_callout(doc, "ปิดหน้าจอได้ไหม", "ปิดหน้าจอได้เมื่ออนุญาตการทำงานเบื้องหลังและยกเว้นการประหยัดแบตเตอรี่แล้ว หากระบบขึ้น ขาดการเชื่อมต่อ ให้เปิดแอป ตรวจ Mobile data และกดเริ่มตรวจตลอดเวลาอีกครั้ง", "FEF3C7", "92400E")
    doc.add_heading("ความหมายของยอดรายเครื่อง", level=2)
    add_bullets(doc, [
        "ใช้งานได้ / โหลดช้า / ใช้ไม่ได้ / ยังไม่ทราบ รวมกันต้องเท่ากับ URL ที่เครื่องนั้นได้รับมอบหมาย",
        "เคสค้างเป็นอีกมิติหนึ่ง ไม่ควรนำไปบวกกับจำนวน URL เพราะเคสนับแยกตามบริษัทและห้อง",
        "กดการ์ดสถานะหรือ ดูรายละเอียด เพื่อกรองรายการด้านล่างตามสถานะ",
    ])
    add_picture(doc, "mobile-result-detail.png", "รายละเอียดผลตรวจบนมือถือ — ใช้ค้นหาและเปิดรายการแต่ละ URL", width=3.05)

    add_title(doc, "7. รายงานรอบวัน KPI และ Telegram")
    add_picture(doc, "report.jpg", "รายงานรอบวัน — สรุปจำนวนเคส แก้แล้ว ค้าง และแยกตามรอบ", width=5.55)
    doc.add_heading("รายงานรอบวัน", level=2)
    add_bullets(doc, [
        "เลือกวันที่และบริษัทก่อนอ่านยอด",
        "ตรวจเคสค้างและเวลาที่เริ่มค้าง รวมทั้งแยก URL หลักกับลิงก์สำรอง",
        "ตรวจสรุปรอบเช้า เย็น กลางคืน และยอดแยกตามเครื่อง/ค่าย",
        "ส่งออก PNG หรือ PDF หลังจากเลือกข้อมูลที่ต้องการแล้ว",
    ], numbered=True)
    doc.add_heading("KPI รายคน", level=2)
    add_bullets(doc, [
        "ใช้ดูจำนวนเคสที่รับ เวลาตอบสนอง เวลาปิดงาน และเคสค้าง",
        "MANAGEMENT อ่านได้อย่างเดียว ส่วน ADMIN_COMPANY ไม่เห็น KPI",
        "ถ้ายอดไม่ตรง ให้ตรวจตัวกรองบริษัท วันที่ และแหล่งตรวจที่เลือก",
    ])
    doc.add_heading("Telegram", level=2)
    add_bullets(doc, [
        "แจ้งเมื่อยืนยันปัญหาจริงตามเกณฑ์ ไม่แจ้งจาก timeout เพียงรอบเดียว",
        "แจ้งชื่อบริษัท ห้อง LINE URL สาเหตุ แหล่งตรวจ เครื่อง และค่าย",
        "เมื่อลิงก์กลับมาใช้งานได้ จะส่งข้อความกลับมาแล้ว แม้ยังช้าจะระบุว่า กลับมาแล้วแต่โหลดช้า",
        "กลุ่มรับแจ้งต้องผูกกับบริษัท/ห้องตามการตั้งค่า เพื่อไม่ส่งข้อมูลข้ามกลุ่ม",
    ])

    add_title(doc, "8. แก้ปัญหาที่พบบ่อยและตรวจรับงาน")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, heading, width in zip(table.rows[0].cells, ("อาการ", "วิธีตรวจและแก้"), (2.15, 4.5)):
        cell.text = heading
        set_cell_width(cell, width)
        set_cell_fill(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    troubleshooting = [
        ("ไม่เห็นเมนู", "ตรวจว่าเข้าสู่ระบบแล้ว บนมือถือแตะปุ่มเมนูมุมซ้ายบน และตรวจบทบาทผู้ใช้"),
        ("เครื่องขึ้นขาดการเชื่อมต่อ", "เปิดแอป ตรวจ Mobile data ยกเว้นประหยัดแบตเตอรี่ และกดเริ่มตรวจตลอดเวลา"),
        ("VPN ไม่แสดงตำแหน่ง", "ตรวจว่า Surfshark ขึ้น Connected และเลือก Thailand/Bangkok, อัปเดต Agent เป็นรุ่น 1.0.6 ขึ้นไป แล้วรอผลรอบใหม่ประมาณ 5 นาที"),
        ("Unauthorized", "QR หมดอายุหรือเครื่องถูกย้าย ให้สร้าง QR ใหม่แล้วผูกอีกครั้ง"),
        ("Cleartext HTTP not permitted", "อัปเดต Agent เป็นรุ่นล่าสุด และตรวจว่า URL ใช้ https เมื่อเว็บไซต์รองรับ"),
        ("UnknownHostException", "ตรวจ DNS/โดเมนผ่านซิม หาก URL หลักถูกบล็อกให้ใส่ลิงก์สำรองที่ใช้งานได้"),
        ("แก้แล้วเคสยังไม่ปิด", "ตรวจว่าอยู่สถานะรอยืนยัน รอผลเครื่องซิมรอบใหม่ และยืนยันว่าลิงก์สำรองถูกบันทึก"),
        ("ยอดไม่ตรง", "ตรวจตัวกรอง วันที่ บริษัท แหล่งตรวจ และแยกจำนวน URL ออกจากจำนวนเคส"),
        ("ดึง Rich Menu ไม่ได้", "บันทึก Channel Access Token ก่อน ตรวจว่าเป็น Messaging API channel และเมนูถูกสร้างผ่าน Messaging API; เมนูจาก OA Manager อาจดึงผ่าน API ไม่ได้"),
    ]
    for symptom, fix in troubleshooting:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, (symptom, fix), (2.15, 4.5)):
            cell.text = value
            set_cell_width(cell, width)
        cells[0].paragraphs[0].runs[0].bold = True

    doc.add_heading("รายการตรวจรับก่อนส่งมอบ", level=2)
    add_bullets(doc, [
        "ทดลองเข้าสู่ระบบทั้งจากลิงก์และ QR",
        "ตรวจเมนูและปุ่มด้วยบัญชีแต่ละบทบาท",
        "เพิ่ม/แก้ไข/พักเฝ้าดูลิงก์ และตรวจว่าประวัติยังอยู่",
        "ตรวจ LINE OA ด้วย token ทดสอบดึง Rich Menu และตรวจว่า URL ซ้ำไม่ถูกสร้างเพิ่ม",
        "ผูกเครื่องซิม ตรวจเมื่อปิดหน้าจอ และทดสอบย้ายเครื่องด้วย QR ใหม่",
        "จำลองปัญหา ตรวจ Telegram เปิดเคส รอยืนยัน และปิดเคส",
        "ตรวจยอด Dashboard เหตุการณ์ รายงานรอบวัน และ KPI ว่าใช้ตัวกรองเดียวกัน",
        "ทดลองส่งออก PNG/PDF และเปิดไฟล์บนโทรศัพท์",
    ])
    add_callout(doc, "ช่องทางช่วยเหลือ", "เมื่อพบปัญหา ให้บันทึกเวลา ชื่อผู้ใช้ บริษัท ห้อง LINE URL รหัสเคส ชื่อเครื่อง และภาพหน้าจอ ก่อนส่งให้ผู้ดูแลระบบ", "EFF6FF", BLUE)

    force_fonts(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
