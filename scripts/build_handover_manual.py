from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "public"
ASSETS = OUT / "manual-assets"
OUT.mkdir(parents=True, exist_ok=True)
ASSETS.mkdir(parents=True, exist_ok=True)

BLUE = "2458E6"
NAVY = "172554"
INK = "1E293B"
MUTED = "64748B"
LIGHT = "F5F7FC"
GREEN = "059669"
AMBER = "D97706"
RED = "DC2626"
FONT_PATH = "/System/Library/Fonts/Supplemental/Tahoma.ttf"
FONT_BOLD_PATH = "/System/Library/Fonts/Supplemental/Tahoma Bold.ttf"


def font(size, bold=False):
    return ImageFont.truetype(FONT_BOLD_PATH if bold else FONT_PATH, size)


def rounded(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text(draw, xy, value, size, fill="#1E293B", bold=False, anchor=None):
    draw.text(xy, value, font=font(size, bold), fill=fill, anchor=anchor)


def pseudo_qr(draw, left, top, size):
    n = 29
    cell = size // n
    draw.rectangle((left, top, left + cell*n, top + cell*n), fill="white")
    def finder(x, y):
        draw.rectangle((left+x*cell, top+y*cell, left+(x+7)*cell, top+(y+7)*cell), fill="black")
        draw.rectangle((left+(x+1)*cell, top+(y+1)*cell, left+(x+6)*cell, top+(y+6)*cell), fill="white")
        draw.rectangle((left+(x+2)*cell, top+(y+2)*cell, left+(x+5)*cell, top+(y+5)*cell), fill="black")
    finder(1,1); finder(21,1); finder(1,21)
    for y in range(n):
        for x in range(n):
            if (x < 9 and y < 9) or (x > 19 and y < 9) or (x < 9 and y > 19):
                continue
            if ((x*17 + y*31 + x*y*3) % 7) < 3:
                draw.rectangle((left+x*cell, top+y*cell, left+(x+1)*cell, top+(y+1)*cell), fill="black")


def make_architecture():
    im = Image.new("RGB", (1600, 900), "#F5F7FC")
    d = ImageDraw.Draw(im)
    text(d, (80, 62), "Flow การตรวจจากซิม TRUE", 48, "#172554", True)
    text(d, (80, 126), "โทรศัพท์บังคับคำขอออกผ่าน Cellular แล้วส่งผลกลับ DomainWatch", 26, "#64748B")
    cards = [
        (90, 260, 430, 600, "1", "Android Agent", "ซิม TRUE\nตรวจทุก 5 นาที\n32 URL พร้อมกัน", "#2458E6"),
        (630, 260, 970, 600, "2", "DomainWatch API", "รับผลแยกตาม URL\nยืนยันผิดปกติ 2 รอบ\nเก็บประวัติเหตุการณ์", "#7C3AED"),
        (1170, 260, 1510, 600, "3", "Dashboard + Telegram", "เขียว = ใช้งานได้\nเหลือง = ช้า\nแดง = ปัญหาจริง", "#059669"),
    ]
    for x1,y1,x2,y2,num,title,body,color in cards:
        rounded(d,(x1,y1,x2,y2),28,"white","#E2E8F0",3)
        rounded(d,(x1+30,y1+30,x1+100,y1+100),18,color)
        text(d,(x1+65,y1+65),num,34,"white",True,"mm")
        text(d,(x1+30,y1+145),title,34,"#1E293B",True)
        for i,line in enumerate(body.split("\n")):
            text(d,(x1+30,y1+215+i*52),line,25,"#475569")
    for a,b in [((430,430),(630,430)),((970,430),(1170,430))]:
        d.line((*a,*b), fill="#94A3B8", width=8)
        d.polygon([(b[0],b[1]),(b[0]-28,b[1]-18),(b[0]-28,b[1]+18)], fill="#94A3B8")
    rounded(d,(90,700,1510,820),24,"#EFF6FF")
    text(d,(120,738),"หลักสำคัญ",24,"#2458E6",True)
    text(d,(300,738),"เว็บช้าไม่ถูกนับเป็นล่ม • ล่ม 2 รอบจึงแจ้ง • กลับมาปกติ 2 รอบจึงปิดเคส",26,"#1E3A8A")
    p=ASSETS/"architecture.png"; im.save(p); return p


def make_admin_mock():
    im=Image.new("RGB",(1500,900),"#F5F7FC"); d=ImageDraw.Draw(im)
    d.rectangle((0,0,260,900),fill="#2448B9")
    text(d,(42,55),"D",48,"white",True); text(d,(105,58),"DomainWatch",28,"white",True)
    for i,label in enumerate(["แดชบอร์ด","เหตุการณ์","เครื่องตรวจเครือข่าย"]):
        y=180+i*90
        if i==2: rounded(d,(24,y-20,236,y+45),15,"white")
        text(d,(48,y),label,22,"#2448B9" if i==2 else "white",i==2)
    text(d,(320,65),"เครื่องตรวจเครือข่ายมือถือ",38,"#172554",True)
    text(d,(320,120),"ตรวจลิงก์จากซิมจริง แยกจากตัวตรวจบน Vercel",22,"#64748B")
    rounded(d,(320,180,1420,355),24,"white","#E2E8F0",2)
    text(d,(355,215),"เพิ่มเครื่องตรวจ TRUE",25,"#1E293B",True)
    text(d,(355,260),"ชื่อเครื่อง",18,"#64748B")
    rounded(d,(355,292,1000,338),10,"#F8FAFC","#CBD5E1",2)
    text(d,(375,304),"เครื่องตรวจ TRUE ห้อง IT",18,"#334155")
    rounded(d,(1040,286,1385,342),12,"#2458E6")
    text(d,(1212,314),"+ สร้างเครื่องและ QR",19,"white",True,"mm")
    rounded(d,(450,390,1290,845),28,"white","#CBD5E1",2)
    text(d,(870,435),"QR ผูกเครื่อง — เครื่องตรวจ TRUE ห้อง IT",25,"#1E293B",True,"mm")
    text(d,(870,477),"ใช้ได้ครั้งเดียว ภายใน 15 นาที",18,"#D97706",True,"mm")
    pseudo_qr(d,690,515,310)
    rounded(d,(1030,555,1240,625),12,"#FEF3C7")
    text(d,(1135,590),"ตัวอย่างเท่านั้น",17,"#92400E",True,"mm")
    text(d,(1135,655),"QR จริงสร้างจากระบบ\nห้ามส่งเข้ากลุ่มสาธารณะ",17,"#64748B",False,"mm")
    p=ASSETS/"admin-qr.png"; im.save(p); return p


def make_phone_mock():
    im=Image.new("RGB",(900,1600),"#EEF2FF"); d=ImageDraw.Draw(im)
    rounded(d,(165,55,735,1545),58,"#111827")
    rounded(d,(185,85,715,1515),44,"#F5F7FC")
    rounded(d,(375,100,525,125),12,"#111827")
    rounded(d,(225,175,325,275),24,"#2458E6"); text(d,(275,225),"D",48,"white",True,"mm")
    text(d,(225,315),"DomainWatch Agent",34,"#1E293B",True)
    text(d,(225,365),"ตรวจหน้าเว็บผ่านซิมมือถือจริง ทุก 5 นาที",19,"#64748B")
    rounded(d,(225,430,675,650),24,"white","#E2E8F0",2)
    text(d,(255,470),"🟢  กำลังตรวจตลอดเวลา",27,"#059669",True)
    text(d,(255,525),"เครื่อง: เครื่องตรวจ TRUE ห้อง IT",18,"#475569")
    text(d,(255,565),"ซิมที่กำหนด: TRUE",18,"#475569")
    text(d,(255,605),"คำขอออกผ่าน Cellular",18,"#475569")
    rounded(d,(225,690,675,840),24,"white","#E2E8F0",2)
    text(d,(255,725),"ผลตรวจรอบล่าสุด",18,"#64748B",True)
    text(d,(255,775),"504 URL • ปกติ 498 • ช้า 6",24,"#1E293B",True)
    rounded(d,(225,885,675,960),14,"#2458E6")
    text(d,(450,922),"เริ่มตรวจตลอดเวลา",22,"white",True,"mm")
    rounded(d,(225,985,675,1060),14,"#E8EEFF")
    text(d,(450,1022),"หยุดตรวจชั่วคราว",22,"#173FAD",True,"mm")
    rounded(d,(225,1085,675,1160),14,"white","#CBD5E1",2)
    text(d,(450,1122),"อนุญาตให้ทำงานเบื้องหลัง",19,"#173FAD",True,"mm")
    text(d,(225,1245),"สถานะที่ต้องเห็นหลังติดตั้ง",20,"#172554",True)
    for i,s in enumerate(["แจ้งเตือน DomainWatch Agent ค้างอยู่","Mobile data เปิด และ Wi‑Fi เปิดได้","ตรวจล่าสุดไม่เกิน 5–7 นาที"]):
        text(d,(245,1295+i*52),"✓ "+s,18,"#334155")
    p=ASSETS/"phone-agent.png"; im.save(p); return p


def make_transfer():
    im=Image.new("RGB",(1600,850),"white"); d=ImageDraw.Draw(im)
    text(d,(80,60),"ย้ายโทรศัพท์โดยไม่ต้องตั้งระบบใหม่",44,"#172554",True)
    steps=[
        (120,"1","เครื่องเดิม","กดหยุดตรวจ\nหรือปิดใช้งาน"),
        (520,"2","หน้า /agents","กด สร้าง QR / ย้ายเครื่อง"),
        (920,"3","เครื่องใหม่","ติดตั้ง APK เดิม\nสแกน QR ใหม่"),
        (1320,"4","ตรวจรับ","รอผลรอบแรก\nสถานะออนไลน์"),
    ]
    for x,n,t,b in steps:
        rounded(d,(x-80,210,x+260,610),26,"#F8FAFC","#CBD5E1",3)
        rounded(d,(x-45,245,x+35,325),20,"#2458E6")
        text(d,(x-5,285),n,34,"white",True,"mm")
        text(d,(x-45,370),t,28,"#1E293B",True)
        for i,line in enumerate(b.split("\n")): text(d,(x-45,420+i*45),line,21,"#64748B")
        if x<1320:
            d.line((x+260,410,x+400,410),fill="#94A3B8",width=7)
            d.polygon([(x+400,410),(x+370,392),(x+370,428)],fill="#94A3B8")
    rounded(d,(80,690,1520,790),20,"#FEF3C7")
    text(d,(110,724),"หมายเหตุ",22,"#92400E",True)
    text(d,(275,724),"QR เป็นสิทธิ์ใช้ครั้งเดียว อายุ 15 นาที และจะยกเลิกสิทธิ์เครื่องเดิมอัตโนมัติเมื่อผูกเครื่องใหม่",22,"#78350F")
    p=ASSETS/"transfer.png"; im.save(p); return p


def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)


def set_margins(section):
    section.top_margin=Cm(1.6); section.bottom_margin=Cm(1.5); section.left_margin=Cm(1.7); section.right_margin=Cm(1.7)


def add_page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run=paragraph.add_run("DomainWatch • ")
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); paragraph._p.append(fld)
    run.font.size=Pt(9); run.font.color.rgb=RGBColor.from_string(MUTED)


def style_doc(doc):
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Tahoma'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Tahoma'); normal._element.rPr.rFonts.set(qn('w:cs'),'Tahoma'); normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(INK)
    normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.08
    for name,size,color in [('Title',32,NAVY),('Heading 1',22,NAVY),('Heading 2',15,BLUE),('Heading 3',12,INK)]:
        s=styles[name]; s.font.name='Tahoma'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Tahoma'); s._element.rPr.rFonts.set(qn('w:cs'),'Tahoma'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
    if 'Caption Thai' not in styles:
        s=styles.add_style('Caption Thai',WD_STYLE_TYPE.PARAGRAPH); s.font.name='Tahoma'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Tahoma'); s._element.rPr.rFonts.set(qn('w:cs'),'Tahoma'); s.font.size=Pt(8.5); s.font.color.rgb=RGBColor.from_string(MUTED); s.paragraph_format.space_after=Pt(8)


def force_thai_fonts(doc):
    def paragraphs_in(parent):
        for p in parent.paragraphs:
            yield p
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from paragraphs_in(cell)
    for p in paragraphs_in(doc):
        for r in p.runs:
            r.font.name='Tahoma'
            rpr=r._element.get_or_add_rPr()
            rfonts=rpr.rFonts
            if rfonts is None:
                rfonts=OxmlElement('w:rFonts'); rpr.insert(0,rfonts)
            for key in ('ascii','hAnsi','eastAsia','cs'):
                rfonts.set(qn('w:'+key),'Tahoma')
            lang=rpr.find(qn('w:lang'))
            if lang is None:
                lang=OxmlElement('w:lang'); rpr.append(lang)
            lang.set(qn('w:val'),'th-TH'); lang.set(qn('w:bidi'),'th-TH')
    for section in doc.sections:
        for part in (section.header,section.footer):
            for p in paragraphs_in(part):
                for r in p.runs:
                    r.font.name='Tahoma'
                    rpr=r._element.get_or_add_rPr(); rfonts=rpr.rFonts
                    if rfonts is None:
                        rfonts=OxmlElement('w:rFonts'); rpr.insert(0,rfonts)
                    for key in ('ascii','hAnsi','eastAsia','cs'):
                        rfonts.set(qn('w:'+key),'Tahoma')


def add_header_footer(section):
    h=section.header.paragraphs[0]; h.text="DOMAINWATCH  |  TRUE MOBILE AGENT"; h.style='Caption Thai'
    add_page_number(section.footer.paragraphs[0])


def add_title(doc,title,subtitle=None):
    doc.add_heading(title,0)
    if subtitle:
        p=doc.add_paragraph(subtitle); p.style='Subtitle'


def add_callout(doc,title,body,color="EFF6FF"):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
    c=t.cell(0,0); set_cell_shading(c,color); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p=c.paragraphs[0]; r=p.add_run(title+"\n"); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); r.font.size=Pt(12)
    r=p.add_run(body); r.font.size=Pt(10)
    doc.add_paragraph()


def add_steps(doc,items):
    table=doc.add_table(rows=0,cols=2); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    for i,(title,body) in enumerate(items,1):
        cells=table.add_row().cells; cells[0].width=Cm(1.2); cells[1].width=Cm(15.5)
        set_cell_shading(cells[0],BLUE); cells[0].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=cells[0].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(str(i)); r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(15)
        p=cells[1].paragraphs[0]; r=p.add_run(title+"\n"); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(INK)
        r=p.add_run(body); r.font.size=Pt(9.5); r.font.color.rgb=RGBColor.from_string(MUTED)
    doc.add_paragraph()


def add_checklist(doc,items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent=Cm(.4); p.add_run("☐ "+item)


def add_page(doc,title,subtitle=None):
    doc.add_page_break(); add_title(doc,title,subtitle)


def build_doc():
    arch=make_architecture(); admin=make_admin_mock(); phone=make_phone_mock(); transfer=make_transfer()
    doc=Document(); style_doc(doc)
    for s in doc.sections: set_margins(s); add_header_footer(s)
    sec=doc.sections[0]; sec.header.is_linked_to_previous=False

    # Cover
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(45); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("D"); r.bold=True; r.font.size=Pt(50); r.font.color.rgb=RGBColor.from_string(BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("DOMAINWATCH"); r.bold=True; r.font.size=Pt(17); r.font.color.rgb=RGBColor.from_string(NAVY)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(55); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("คู่มือส่งมอบ\nเครื่องตรวจเครือข่ายมือถือ TRUE"); r.bold=True; r.font.size=Pt(29); r.font.color.rgb=RGBColor.from_string(NAVY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("ติดตั้ง • ผูกเครื่องด้วย QR • ตรวจตลอดเวลา • ย้ายเครื่องได้"); r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string(MUTED)
    doc.add_picture(str(phone),width=Cm(7.0)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("เวอร์ชันเอกสาร 1.0  |  แอป Android 1.0.0  |  26 สิงหาคม 2569"); r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(MUTED)

    add_page(doc,"เริ่มใช้งานแบบเร็ว","ใช้เวลาโดยประมาณ 10–15 นาที เมื่อโทรศัพท์มีซิม TRUE และอินเทอร์เน็ตพร้อม")
    add_callout(doc,"สิ่งที่ได้รับ","APK ที่ลงนามแล้ว, หน้าเมนูเครื่องตรวจเครือข่าย, QR ใช้ครั้งเดียว, ประวัติผลตรวจจากซิม และการแจ้งเตือน Telegram แยกตาม URL/ห้อง LINE")
    add_steps(doc,[
        ("ดาวน์โหลดและติดตั้ง APK","เข้า DomainWatch ด้วยสิทธิ์ ADMIN → เมนูเครื่องตรวจเครือข่าย → ดาวน์โหลด APK"),
        ("สร้างเครื่อง TRUE","ตั้งชื่อ เช่น “เครื่องตรวจ TRUE ห้อง IT” แล้วกด + สร้างเครื่องและ QR"),
        ("สแกน QR","ใช้กล้องโทรศัพท์สแกน QR ภายใน 15 นาที แล้วกดเปิดด้วย DomainWatch Agent"),
        ("อนุญาตทำงานเบื้องหลัง","อนุญาตแจ้งเตือน และกด “อนุญาตให้ทำงานเบื้องหลัง”"),
        ("ตรวจรับ","รอรอบแรก 5–7 นาที หน้า /agents ต้องเห็นออนไลน์และมีผล URL ล่าสุด"),
    ])
    doc.add_heading("ของที่ต้องเตรียม",1)
    add_checklist(doc,["โทรศัพท์ Android 8 ขึ้นไป พร้อมที่ชาร์จถาวร","ซิม TRUE เปิดใช้งาน Mobile data และมีแพ็กเกจอินเทอร์เน็ต","บัญชี DomainWatch บทบาท ADMIN","Wi‑Fi เปิดไว้ได้ แต่แอปจะบังคับการตรวจผ่าน Cellular","ตำแหน่งวางเครื่องมีสัญญาณ TRUE เสถียรและระบายความร้อนได้"])

    add_page(doc,"ระบบทำงานอย่างไร","ตัวตรวจจากซิมจริงทำงานแยกจากตัวตรวจบน Vercel เพื่อยืนยันปัญหาที่ผู้ใช้เครือข่ายมือถือพบจริง")
    doc.add_picture(str(arch),width=Cm(17.2)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph("ภาพที่ 1 — Flow การตรวจจากโทรศัพท์ซิม TRUE ไปยัง DomainWatch และ Telegram"); p.style='Caption Thai'; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("กติกาความแม่นยำ",1)
    add_steps(doc,[
        ("ตรวจทุก 5 นาที","ดึงรายการ URL จากระบบและตรวจพร้อมกันสูงสุด 32 งาน"),
        ("ช้าไม่ใช่ล่ม","ตอบกลับได้แต่เกินเกณฑ์แสดงเป็นสีเหลือง โดยไม่เปิดเหตุการณ์ล่ม"),
        ("ยืนยัน 2 รอบ","ต้องผิดปกติต่อเนื่อง 2 รอบจึงเปิดเคสและส่ง Telegram ลด false alarm"),
        ("กลับมา 2 รอบ","ต้องปกติต่อเนื่อง 2 รอบจึงปิดเคสและแจ้งว่ากลับมาแล้ว แม้ยังช้า"),
    ])

    add_page(doc,"ตั้งค่าเครื่องจากหน้าแอดมิน","เมนูนี้เห็นได้เฉพาะ ADMIN และทุก API มีการตรวจสิทธิ์ฝั่ง server")
    doc.add_picture(str(admin),width=Cm(17.4)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph("ภาพที่ 2 — ตัวอย่างหน้า /agents และ QR ตัวอย่าง (ไม่ใช่ QR ใช้งานจริง)"); p.style='Caption Thai'; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_steps(doc,[
        ("เข้าสู่ระบบ","เปิด https://domain-watch-app-sandy.vercel.app/ และล็อกอินด้วย ADMIN"),
        ("เปิดเมนูเครื่องตรวจเครือข่าย","กดดาวน์โหลด APK ก่อน แล้วติดตั้งบนโทรศัพท์"),
        ("กรอกชื่อเครื่อง","ตั้งชื่อให้รู้ตำแหน่งและค่าย เช่น เครื่องตรวจ TRUE ห้อง IT"),
        ("สร้าง QR","QR ใช้ครั้งเดียวและหมดอายุใน 15 นาที ห้ามส่งเข้ากลุ่มสาธารณะ"),
    ])

    add_page(doc,"ติดตั้ง APK บนโทรศัพท์ Android","ไม่ต้องใช้ Termux, Termux:Boot หรือ Termux:API")
    add_steps(doc,[
        ("ดาวน์โหลด APK","จากหน้า /agents กด ดาวน์โหลด APK หรือใช้ไฟล์ DomainWatch-Agent-v1.0.0.apk ในชุดส่งมอบ"),
        ("อนุญาตติดตั้งแอปไม่รู้จัก","หาก Android ถาม ให้เปิดอนุญาตเฉพาะเบราว์เซอร์/แอปไฟล์ที่ใช้ติดตั้ง แล้วปิดสิทธิ์นี้ภายหลังได้"),
        ("ติดตั้ง","กด ติดตั้ง และเปิด DomainWatch Agent เมื่อเสร็จ"),
        ("อนุญาตแจ้งเตือน","ต้องอนุญาตเพื่อให้บริการเบื้องหลังแสดงสถานะต่อเนื่อง"),
    ])
    add_callout(doc,"ตรวจลายเซ็นไฟล์","SHA‑256 ของ APK รุ่น 1.0.0: c2c43a431e7f839947586c644d2f48d860f6c42149fdfde248d09ccac17e4e6b","F0FDF4")
    doc.add_heading("ถ้าติดตั้งไม่ได้",1)
    add_checklist(doc,["ลบ APK ที่ดาวน์โหลดไม่ครบแล้วดาวน์โหลดใหม่","ตรวจพื้นที่ว่างอย่างน้อย 300 MB","หากมีแอปรุ่นทดสอบคนละลายเซ็น ให้ถอนรุ่นทดสอบก่อน","อย่าแชร์ APK ผ่านแอปที่บีบอัดหรือเปลี่ยนชื่อไฟล์"])

    add_page(doc,"ผูกเครื่องด้วย QR และเปิดตรวจตลอดเวลา","ขั้นตอนนี้ต้องให้ Mobile data ของ TRUE เปิดอยู่")
    doc.add_picture(str(phone),height=Cm(14.0)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_steps(doc,[
        ("สแกน QR ด้วยกล้อง","กล้องจะเปิดลิงก์ https ของ DomainWatch แล้วส่งต่อเข้าแอปด้วย domainwatch-agent://"),
        ("เลือกเปิดด้วย DomainWatch Agent","แอปจะยืนยัน QR ผ่านเครือข่าย Cellular และเก็บ token แบบเข้ารหัสใน Android Keystore"),
        ("กดเริ่มตรวจตลอดเวลา","สถานะต้องเป็น “กำลังตรวจตลอดเวลา” และมีแจ้งเตือนค้างอยู่"),
        ("ยกเว้นการประหยัดแบตเตอรี่","กดปุ่มอนุญาตทำงานเบื้องหลัง แล้วเลือก อนุญาต/ไม่จำกัด"),
    ])

    # The preceding QR page fills the sheet exactly; adding another explicit page
    # break makes LibreOffice emit a blank page. Let normal pagination place this
    # heading at the top of the next page instead.
    add_title(doc,"ตรวจรับหลังติดตั้ง","อย่าถือว่าเสร็จจนกว่าจะเห็นผลรอบแรกจากซิมจริง")
    table=doc.add_table(rows=1,cols=3); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
    headers=["จุดตรวจ","ต้องเห็น","ถ้าไม่ผ่าน"]
    for i,h in enumerate(headers): set_cell_shading(table.rows[0].cells[i],NAVY); r=table.rows[0].cells[i].paragraphs[0].add_run(h); r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
    rows=[
        ("โทรศัพท์","🟢 กำลังตรวจตลอดเวลา","เปิด Mobile data และกดเริ่มตรวจ"),
        ("หน้า /agents","สถานะออนไลน์ ภายใน 5–7 นาที","ตรวจแบตเตอรี่/สัญญาณ/เวลาเครื่อง"),
        ("ผล URL","มี HTTP code และ response time","รออีกรอบ หรือเปิดแอปดู error"),
        ("Telegram","ไม่มีข้อความล่มจากเว็บที่เพียงช้า","ดู failure streak ต้องครบ 2/2"),
        ("รีสตาร์ตเครื่อง","บริการกลับมาทำงานอัตโนมัติ","เปิดแอปหนึ่งครั้งและอนุญาตพื้นหลัง"),
    ]
    for row in rows:
        cells=table.add_row().cells
        for i,val in enumerate(row): cells[i].text=val; cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()
    add_callout(doc,"เกณฑ์ออนไลน์","หน้า /agents ถือว่าเครื่องออนไลน์เมื่อมีผลตรวจล่าสุดไม่เกินประมาณ 12 นาที การตรวจปกติจะส่งทุก 5 นาที")
    doc.add_heading("ทดสอบจริงอย่างปลอดภัย",1)
    add_checklist(doc,["ปิด Mobile data 10–12 นาทีเพื่อทดสอบสถานะขาดการเชื่อมต่อ แล้วเปิดกลับ","ใช้ URL ทดสอบเฉพาะที่ทีมควบคุมได้ ห้ามทำให้เว็บจริงล่ม","ตรวจว่า Telegram แจ้งล่มเมื่อผิดปกติ 2 รอบและแจ้งกลับเมื่อปกติ 2 รอบ"])

    add_page(doc,"การแจ้งเตือนและประวัติเหตุการณ์","เหตุการณ์จากซิมถูกเก็บแยกและอ้างอิง URL/บริษัท/ห้อง LINE ชัดเจน")
    table=doc.add_table(rows=1,cols=4); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(["สถานะ","เงื่อนไข","Dashboard","Telegram"]): set_cell_shading(table.rows[0].cells[i],NAVY); rr=table.rows[0].cells[i].paragraphs[0].add_run(h); rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)
    data=[
        ("🟢 ใช้งานได้","ตอบกลับในเกณฑ์","UP","ไม่แจ้งซ้ำ"),
        ("🟡 โหลดช้า","ตอบได้แต่เกินเกณฑ์","SLOW","ไม่เปิดเคสล่ม"),
        ("🔴 ใช้ไม่ได้","ผิดปกติ 2 รอบติด","DOWN + เปิดเคส","แจ้งล่มพร้อม URL/ห้อง"),
        ("🟢 กลับมาแล้ว","ปกติ 2 รอบติด","ปิดเคส","แจ้งกลับมา แม้ยังช้า"),
    ]
    for row in data:
        cells=table.add_row().cells
        for i,v in enumerate(row): cells[i].text=v
    doc.add_paragraph()
    add_callout(doc,"กรณี URL เดียวกันหลายห้อง LINE","ระบบเปิดเหตุการณ์และแจ้งแยกตามลิงก์/ห้อง LINE จึงทราบว่ารายการใดได้รับผลกระทบ แม้ปลายทาง URL เหมือนกัน")
    doc.add_heading("ข้อมูลที่ต้องมีในประวัติ",1)
    add_checklist(doc,["รหัสเคสและชื่อหน้า","บริษัทและห้อง LINE","URL ที่ตรวจจริง","เครือข่าย TRUE และเครื่องตรวจ","เวลาเริ่มค้าง/เวลาปิดเคส","HTTP code, response time และสาเหตุล่าสุด"])

    add_page(doc,"ย้ายโทรศัพท์หรือเปลี่ยนซิม","ไม่ต้องแก้โค้ด ไม่ต้องสร้างระบบใหม่ และไม่ต้องใช้ QR เดิม")
    doc.add_picture(str(transfer),width=Cm(17.2)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph("ภาพที่ 4 — Flow ย้ายเครื่องด้วย QR ใหม่"); p.style='Caption Thai'; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_steps(doc,[
        ("เตรียมเครื่องใหม่","ใส่ซิม TRUE เปิด Mobile data ติดตั้ง APK รุ่นเดิมหรือใหม่กว่าที่ลงนามด้วย key เดิม"),
        ("สร้าง QR ใหม่","หน้า /agents เลือกเครื่องเดิม แล้วกด สร้าง QR / ย้ายเครื่อง"),
        ("สแกนบนเครื่องใหม่","เมื่อผูกสำเร็จ token เครื่องเดิมจะใช้ต่อไม่ได้"),
        ("ตรวจรับและเก็บเครื่องเดิม","รอออนไลน์/ผลรอบแรก จากนั้นถอนแอปหรือเก็บเป็นอะไหล่"),
    ])

    add_page(doc,"ความปลอดภัยและชุดส่งมอบ","เก็บกุญแจเซ็นแอปแยกจาก GitHub และจำกัดสิทธิ์ผู้ดูแล")
    add_callout(doc,"ห้ามอัปโหลดขึ้น GitHub","ไฟล์ domainwatch-agent-release.jks และรหัสผ่านใน SIGNING-KEY-README.txt เป็นความลับสำหรับเซ็นอัปเดต APK เท่านั้น","FEF2F2")
    doc.add_heading("รายการไฟล์ส่งมอบ",1)
    table=doc.add_table(rows=1,cols=3); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(["ไฟล์/ตำแหน่ง","วัตถุประสงค์","การเก็บ"]): set_cell_shading(table.rows[0].cells[i],NAVY); rr=table.rows[0].cells[i].paragraphs[0].add_run(h); rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)
    rows=[
        ("deliverables/public/DomainWatch-Agent-v1.0.0.apk","ติดตั้งบนโทรศัพท์","แชร์ให้ทีมติดตั้งได้"),
        ("android-agent/","ซอร์ส Android สำหรับพัฒนาต่อ","GitHub"),
        ("deliverables/private/*.jks","กุญแจเซ็นอัปเดต APK","Password manager/สื่อเข้ารหัส"),
        ("deliverables/private/SIGNING-KEY-README.txt","ชื่อ alias และขั้นตอนเซ็น","เก็บคู่กับ key แบบจำกัดสิทธิ์"),
        ("คู่มือ DOCX/PDF นี้","ติดตั้ง/ตรวจรับ/ย้ายเครื่อง","คลังเอกสารทีม"),
    ]
    for row in rows:
        cells=table.add_row().cells
        for i,v in enumerate(row): cells[i].text=v
    doc.add_paragraph()
    doc.add_heading("สิทธิ์ระบบ",1)
    add_checklist(doc,["สร้าง/ย้าย/ปิด Mobile Agent: ADMIN เท่านั้น","QR ใช้ครั้งเดียว อายุ 15 นาที และเก็บเฉพาะ hash ฝั่ง server","token โทรศัพท์เก็บแบบเข้ารหัสด้วย Android Keystore","เปลี่ยนเครื่องแล้ว token เครื่องเดิมถูกยกเลิก","อย่าใส่ URL enrollment หรือ QR จริงในคู่มือ/แชตสาธารณะ"])

    add_page(doc,"Troubleshooting และเช็กลิสต์ส่งมอบ","ใช้หน้านี้เป็นใบตรวจรับก่อนปิดงาน")
    issues=[
        ("QR เปิดเว็บแต่ไม่เข้าแอป","ติดตั้ง APK ให้เสร็จ เปิดแอปหนึ่งครั้ง แล้วสแกน QR ใหม่"),
        ("QR หมดอายุ","กด สร้าง QR / ย้ายเครื่อง ใหม่ ใช้ภายใน 15 นาที"),
        ("ผูกไม่สำเร็จ","เปิด Mobile data TRUE, ปิด VPN/Data Saver ชั่วคราว และตั้งวันเวลาอัตโนมัติ"),
        ("เครื่องออฟไลน์หลังดับจอ","ตั้งแบตเตอรี่เป็นไม่จำกัด, อนุญาต autostart และเสียบชาร์จ"),
        ("เห็นเว็บล่มทั้งที่เปิดได้","ตรวจว่า agent ใช้ Cellular จริง, ดู HTTP/redirect และรอการยืนยัน 2 รอบ"),
        ("เว็บกลับมาแต่ช้า","ระบบควรปิดเคสหลังปกติ 2 รอบและแจ้งกลับมาแล้ว พร้อมคงสถานะ SLOW"),
        ("ต้องอัปเดต APK","ต้องเซ็นด้วยไฟล์ .jks เดิมและ versionCode สูงกว่าเดิม"),
    ]
    table=doc.add_table(rows=1,cols=2); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(["อาการ","วิธีแก้"]): set_cell_shading(table.rows[0].cells[i],NAVY); rr=table.rows[0].cells[i].paragraphs[0].add_run(h); rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)
    for a,b in issues:
        c=table.add_row().cells;c[0].text=a;c[1].text=b
    doc.add_paragraph(); doc.add_heading("ลงนามตรวจรับ",1)
    add_checklist(doc,["ติดตั้ง APK และตรวจ SHA‑256 แล้ว","สร้างเครื่อง TRUE และผูก QR สำเร็จ","เห็นสถานะออนไลน์และผลตรวจรอบแรก","ทดสอบรีสตาร์ต/ดับจอแล้ว agent ยังทำงาน","ยืนยัน Telegram ล่ม/กลับมาและประวัติในระบบ","สำรอง signing key ในที่ปลอดภัยอย่างน้อย 2 ชุด","ผู้รับมอบได้รับ DOCX, PDF, APK และตำแหน่ง private key"])
    p=doc.add_paragraph("\nผู้ส่งมอบ ____________________    วันที่ ____________     ผู้รับมอบ ____________________    วันที่ ____________")
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

    force_thai_fonts(doc)
    path=OUT/"DomainWatch-TRUE-Agent-คู่มือส่งมอบ-v1.0.docx"
    doc.save(path)
    print(path)


if __name__ == "__main__":
    build_doc()
