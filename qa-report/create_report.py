from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path('/Users/a/Desktop/DomainWatch-App')
OUT = ROOT / 'qa-report' / 'DomainWatch-System-QA-Report-2026-08-29.docx'
EVIDENCE = ROOT / 'qa-report' / 'evidence'

BLUE = '2446B8'
BLUE_LIGHT = 'EEF3FF'
INK = '172033'
MUTED = '64748B'
GREEN = '087F5B'
GREEN_LIGHT = 'E9F8F1'
AMBER = 'B45309'
AMBER_LIGHT = 'FFF7DF'
RED = 'C92A2A'
RED_LIGHT = 'FFF0F0'
SLATE_LIGHT = 'F3F6FA'
WHITE = 'FFFFFF'


def set_cell_fill(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), color)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, enabled=True):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn('w:keepNext'))
    if enabled and keep is None:
        p_pr.append(OxmlElement('w:keepNext'))


def set_repeat_header_font(run, font='Kanit'):
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    run._element.rPr.rFonts.set(qn('w:cs'), font)


def shade_paragraph(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    p_pr.append(shd)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instruction
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = '1'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, text, end])


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.add_run(text)
    set_keep_with_next(p)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    return p


def set_row_cant_split(row):
    """Keep a compact evidence/finding row together across page breaks."""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn('w:cantSplit')) is None:
        tr_pr.append(OxmlElement('w:cantSplit'))


def add_image(doc: Document, filename: str, caption: str, width=6.9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_keep_with_next(p)
    p.add_run().add_picture(str(EVIDENCE / filename), width=Inches(width))
    add_caption(doc, caption)


def add_status_table(doc: Document, rows):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    widths = [Cm(1.6), Cm(5.2), Cm(2.7), Cm(8.2)]
    hdr = table.rows[0]
    for idx, (txt, w) in enumerate(zip(['ลำดับ', 'หัวข้อทดสอบ', 'ผล', 'หลักฐาน / หมายเหตุ'], widths)):
        hdr.cells[idx].width = w
        set_cell_fill(hdr.cells[idx], BLUE)
        set_cell_margins(hdr.cells[idx])
        p = hdr.cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
    set_repeat_table_header(hdr)
    palette = {
        'ผ่าน': (GREEN_LIGHT, GREEN),
        'ผ่านบางส่วน': (AMBER_LIGHT, AMBER),
        'ไม่ผ่าน': (RED_LIGHT, RED),
        'ไม่ส่งจริง': (SLATE_LIGHT, MUTED),
    }
    for no, topic, status, note in rows:
        cells = table.add_row().cells
        values = [str(no), topic, status, note]
        for idx, value in enumerate(values):
            set_cell_margins(cells[idx], top=65, bottom=65)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            r.font.size = Pt(8.2)
        fill, color = palette.get(status, (SLATE_LIGHT, MUTED))
        set_cell_fill(cells[2], fill)
        cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(color)
        cells[2].paragraphs[0].runs[0].bold = True
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.4)
section.left_margin = Cm(1.55)
section.right_margin = Cm(1.55)
section.header_distance = Cm(0.65)
section.footer_distance = Cm(0.65)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Kanit'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Kanit')
normal._element.rPr.rFonts.set(qn('w:cs'), 'Kanit')
normal.font.size = Pt(9.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.08

for style_name, size, color in [('Title', 29, INK), ('Heading 1', 17, BLUE), ('Heading 2', 12.5, INK), ('Heading 3', 10.5, INK)]:
    s = styles[style_name]
    s.font.name = 'Kanit'
    s._element.rPr.rFonts.set(qn('w:eastAsia'), 'Kanit')
    s._element.rPr.rFonts.set(qn('w:cs'), 'Kanit')
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(7)
    s.paragraph_format.space_after = Pt(5)

# Memo masthead header.
header = section.header
ht = header.add_table(rows=1, cols=2, width=Cm(17.9))
ht.alignment = WD_TABLE_ALIGNMENT.CENTER
ht.rows[0].cells[0].width = Cm(11.8)
ht.rows[0].cells[1].width = Cm(6.1)
left = ht.rows[0].cells[0].paragraphs[0]
left.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = left.add_run('DOMAINWATCH  /  SYSTEM QA')
r.bold = True
r.font.size = Pt(8.5)
r.font.color.rgb = RGBColor.from_string(BLUE)
right = ht.rows[0].cells[1].paragraphs[0]
right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = right.add_run('PRODUCTION • 29 AUG 2026')
r.font.size = Pt(8)
r.font.color.rgb = RGBColor.from_string(MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run('DomainWatch QA Report  •  ')
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor.from_string(MUTED)
add_field(fp, 'PAGE')

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(26)
p.add_run('DOMAINWATCH').bold = True
p.runs[0].font.size = Pt(11)
p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
title = doc.add_paragraph(style='Title')
title.add_run('รายงานผลทดสอบระบบ\nProduction QA')
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_before = Pt(5)
sr = subtitle.add_run('ตรวจครบทั้งระบบกลาง เครื่องตรวจซิม รายงาน สิทธิ์ผู้ใช้ และมุมมองมือถือ')
sr.font.size = Pt(13)
sr.font.color.rgb = RGBColor.from_string(MUTED)

doc.add_paragraph('')
meta = doc.add_table(rows=5, cols=2)
meta.style = 'Table Grid'
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
meta_rows = [
    ('ระบบที่ทดสอบ', 'https://domain-watch-app-sandy.vercel.app/'),
    ('เวอร์ชันโค้ด', 'main / 28fa9d1 — Improve mobile result drill-down'),
    ('ช่วงเก็บหลักฐาน', '29 สิงหาคม 2569 เวลา 09:08–09:24 น. (Asia/Bangkok)'),
    ('ผู้ทดสอบ', 'Codex QA — ใช้บัญชี Test / ADMIN บนระบบจริง'),
    ('รูปแบบทดสอบ', 'Live UI + Responsive + Code review + Automated tests + Production build'),
]
for row, (k, v) in zip(meta.rows, meta_rows):
    set_cell_fill(row.cells[0], BLUE_LIGHT)
    set_cell_margins(row.cells[0], top=110, bottom=110)
    set_cell_margins(row.cells[1], top=110, bottom=110)
    rr = row.cells[0].paragraphs[0].add_run(k)
    rr.bold = True
    rr.font.color.rgb = RGBColor.from_string(BLUE)
    row.cells[1].paragraphs[0].add_run(v)

doc.add_paragraph('')
banner = doc.add_paragraph()
shade_paragraph(banner, BLUE)
banner.paragraph_format.left_indent = Cm(0.3)
banner.paragraph_format.right_indent = Cm(0.3)
banner.paragraph_format.space_before = Pt(10)
banner.paragraph_format.space_after = Pt(8)
br = banner.add_run('สรุป: ระบบแกนหลักทำงานและ build ผ่าน แต่ “ยอดรวมเหตุการณ์” และคำอธิบายสถานะซิมยังทำให้ผู้ใช้ตีความผิด ต้องปรับก่อนถือว่ายอดทุกหน้าตรงกันสมบูรณ์')
br.bold = True
br.font.color.rgb = RGBColor(255, 255, 255)
br.font.size = Pt(11)

doc.add_page_break()

add_heading(doc, '1. Executive summary', 1)
summary = doc.add_table(rows=1, cols=4)
summary.style = 'Table Grid'
summary.alignment = WD_TABLE_ALIGNMENT.CENTER
cards = [
    ('34/34', 'Automated tests', GREEN_LIGHT, GREEN),
    ('PASS', 'Production build', GREEN_LIGHT, GREEN),
    ('651', 'Master Data ทั้งหมด', BLUE_LIGHT, BLUE),
    ('8 จุด', 'ควรปรับปรุง', AMBER_LIGHT, AMBER),
]
for i, (value, label, fill, color) in enumerate(cards):
    c = summary.rows[0].cells[i]
    set_cell_fill(c, fill)
    set_cell_margins(c, top=150, bottom=150)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rv = p.add_run(value + '\n')
    rv.bold = True
    rv.font.size = Pt(18)
    rv.font.color.rgb = RGBColor.from_string(color)
    rl = p.add_run(label)
    rl.font.size = Pt(8)
    rl.font.color.rgb = RGBColor.from_string(MUTED)

doc.add_paragraph('')
for text in [
    'แกนตรวจระบบกลาง การยืนยันล่ม 2 รอบ การกู้คืน 2 รอบ การจำแนก timeout/WAF/redirect และสิทธิ์ 6 บทบาท ผ่านชุดทดสอบอัตโนมัติทั้งหมด 34 รายการ',
    'ยอด Dashboard ณ ภาพหลักฐานครบสมการ: 651 = เว็บไซต์เฝ้าดู 496 + ลิงก์ LINE 155 และ 496 = ใช้ได้ 481 + โหลดช้า/ยังยืนยันไม่ได้ 15 + ใช้ไม่ได้ 0',
    'จุดที่ต้องแก้เร่งด่วนที่สุดคือหน้าเหตุการณ์แสดง “ประวัติทั้งหมด 236” จากข้อมูลที่โหลดสูงสุด 200 เคสต่อประเภท ไม่ใช่จำนวนจริงในฐานข้อมูล ขณะที่ KPI ระบบกลางนับได้ 764 เคส',
    'เครื่องตรวจซิมพบ 10 URL ใช้ไม่ได้ แต่เคสค้างเป็น 0 ได้ตามกติกาเมื่อมีลิงก์สำรองใช้งานได้ อย่างไรก็ตาม UI ควรบอกชัดว่า “ลิงก์หลักมีปัญหา — กำลังใช้ลิงก์สำรอง” เพื่อไม่ให้ดูขัดแย้ง',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)

add_heading(doc, 'คำตัดสิน QA', 2)
decision = doc.add_paragraph()
shade_paragraph(decision, AMBER_LIGHT)
decision.paragraph_format.left_indent = Cm(0.25)
decision.paragraph_format.right_indent = Cm(0.25)
dr = decision.add_run('พร้อมใช้งานแบบมีเงื่อนไข (Conditional Pass) — ฟังก์ชันหลักทำงาน แต่ยังไม่ควรใช้ยอดหน้าเหตุการณ์เป็นยอดรวมทางบริหารจนกว่าจะแก้ count/pagination และควรปรับคำอธิบายสถานะซิมให้แยก URL จริง, รายการตามห้อง และลิงก์สำรอง')
dr.bold = True
dr.font.color.rgb = RGBColor.from_string(AMBER)

add_heading(doc, '2. ยอดจริงที่ยืนยันจากระบบ Production', 1)
actual = doc.add_table(rows=1, cols=4)
actual.style = 'Table Grid'
actual.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['หน้าจอ / เวลา', 'ยอดที่เห็นจริง', 'ตรวจสมการ', 'ผล']
for i, h in enumerate(headers):
    set_cell_fill(actual.rows[0].cells[i], BLUE)
    p = actual.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(8.5)
rows = [
    ('Dashboard 09:11', 'ทั้งหมด 651; เฝ้าดู 496; LINE 155; UP 481; SLOW 15; DOWN 0', '496+155=651; 481+15+0=496', 'ตรง'),
    ('บริษัท 7M', 'ทั้งหมด 68; เฝ้าดู 58; LINE 10; UP 58', '58+10=68', 'ตรง'),
    ('บริษัททั้งหมด', '9 บริษัท; 174 ห้อง; ลิงก์ 651', '112+63+86+32+61+44+138+47+68=651', 'ตรง'),
    ('เครื่องตรวจ TRUE 09:08', 'URL ไม่ซ้ำ 136; UP 97; SLOW 28; DOWN 10; UNKNOWN 1', '97+28+10+1=136', 'ตรง'),
    ('รายงานรอบวัน 09:18', 'เฝ้าดู 496; UP 495; SLOW 1; DOWN 0; เคสวัน 0', '495+1+0=496', 'ตรง'),
    ('KPI', 'เหตุการณ์ระบบกลาง 764; ปิด 764; ผู้ใช้จัดการ 3', 'ยอด KPI มาจาก Incident ทั้งตาราง', 'ตรงตามโค้ด'),
    ('หน้าเหตุการณ์', 'ประวัติทั้งหมด 236; เปิดค้าง 0', 'เป็น 200 ระบบกลาง + 36 ซิมที่โหลดมา', 'ไม่ใช่ยอดจริง'),
]
for row_data in rows:
    cells = actual.add_row().cells
    for i, val in enumerate(row_data):
        set_cell_margins(cells[i], top=65, bottom=65)
        rr = cells[i].paragraphs[0].add_run(val)
        rr.font.size = Pt(7.8)
    if row_data[3] == 'ไม่ใช่ยอดจริง':
        set_cell_fill(cells[3], RED_LIGHT)
        cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(RED)
        cells[3].paragraphs[0].runs[0].bold = True
    else:
        set_cell_fill(cells[3], GREEN_LIGHT)
        cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(GREEN)

doc.add_page_break()

add_heading(doc, '3. ขอบเขตและผลทดสอบรายฟังก์ชัน', 1)
test_rows = [
    (1, 'Login/session และ server guard', 'ผ่าน', 'ชุดทดสอบสิทธิ์และการเข้า route ผ่าน'),
    (2, 'Dashboard ยอดรวมและสถานะ', 'ผ่าน', 'ยอด 651/496/155 และสถานะรวมลงตัว'),
    (3, 'ตัวกรองบริษัท', 'ผ่าน', '7M เปลี่ยนยอดเป็น 68/58/10 ถูกต้อง'),
    (4, 'บริษัท / ห้อง LINE / drill-down', 'ผ่าน', '9 บริษัท; เปิดรายละเอียดห้องและลิงก์ได้'),
    (5, 'Master Data ระบบกลาง', 'ผ่าน', 'พบ 651/651 และตัวกรองทำงาน'),
    (6, 'Master Data เครือข่ายซิม', 'ผ่านบางส่วน', '21 รายการใช้ไม่ได้ แต่เป็นรายการตามห้อง ไม่ใช่ 21 URL จริง'),
    (7, 'เหตุการณ์เปิดค้าง', 'ผ่าน', 'เปิดค้าง 0 สอดคล้องสถานะเคสปัจจุบัน'),
    (8, 'ประวัติเหตุการณ์', 'ไม่ผ่าน', 'ยอด 236 เป็นจำนวนที่โหลด ไม่ใช่ total count'),
    (9, 'Deep link จากเครื่องตรวจไปเคส', 'ผ่าน', 'เปิด modal เคสที่ระบุและแสดงลิงก์สำรองได้'),
    (10, 'แก้ลิงก์จากหน้าเหตุการณ์', 'ผ่านบางส่วน', 'แบบฟอร์มมีข้อมูลครบ; ไม่กดบันทึกบน Production เพื่อไม่เปลี่ยนข้อมูลจริง'),
    (11, 'รายงานรอบวัน', 'ผ่านบางส่วน', 'ยอดระบบกลางตรง; ยอดซิมนับเฉพาะ URL หลัก ไม่รวม backup แต่ป้ายยังไม่ชัด'),
    (12, 'KPI รายคน', 'ผ่านบางส่วน', 'ยอด Incident จริง 764 แต่มีเพียง 3 เคสที่มีผู้รับผิดชอบ KPI'),
    (13, 'เครื่องตรวจซิม/สถานะออนไลน์', 'ผ่าน', 'เครื่อง TRUE-H ออนไลน์ รุ่น/แอป/เวลาล่าสุดแสดงครบ'),
    (14, 'ผลตรวจแยก UP/SLOW/DOWN/UNKNOWN', 'ผ่าน', 'รวม 136 URL ไม่ซ้ำและกด drill-down ได้'),
    (15, 'ยืนยันล่ม/กู้คืน 2 รอบ', 'ผ่าน', 'logic ผ่าน unit tests; UI แสดง streak เกิน 2 แบบ 105/2 ซึ่งชวนสับสน'),
    (16, 'Redirect และลิงก์สำรอง', 'ผ่าน', 'ทดสอบ logic และเห็นเคสปิดผ่านลิงก์สำรองจริง'),
    (17, 'สิทธิ์ 6 บทบาท', 'ผ่าน', 'ADMIN, LEAD, COMPANY, IT, MANAGEMENT, SITE_STAFF ผ่าน 6 tests'),
    (18, 'มือถือ: เมนู/แดชบอร์ด/เครื่องตรวจ', 'ผ่านบางส่วน', 'ใช้งานได้ แต่ปุ่มคู่มือลอยทับ modal บางส่วน'),
    (19, 'ไฟล์ APK', 'ผ่านบางส่วน', 'ลิงก์ชี้รุ่น 1.0.4; browser automation ไม่ได้รับ download event'),
    (20, 'Telegram และ Cron', 'ไม่ส่งจริง', 'ตรวจ message builder/route ด้วย tests; ไม่ส่งเข้ากลุ่มจริงเพื่อไม่รบกวนงาน'),
    (21, 'Export PNG / Print-PDF', 'ผ่านบางส่วน', 'ปุ่มแสดง; automation ไม่ยืนยันไฟล์ดาวน์โหลด PNG'),
    (22, 'Build / TypeScript / ESLint', 'ผ่าน', 'next build, tsc และ lint ผ่านไม่มี error'),
]
add_status_table(doc, test_rows)

doc.add_page_break()

add_heading(doc, '4. ประเด็นที่พบและข้อเสนอปรับปรุง', 1)
findings = [
    ('P1', 'ยอด “ประวัติทั้งหมด” ไม่ใช่ยอดจริง', 'หน้า incidents ใช้ findMany(take: 200) ทั้ง Incident และ NetworkIncident แล้วใช้ array.length แสดงยอด 236 ขณะที่ KPI ดึง Incident ทั้งตารางได้ 764', 'เพิ่ม prisma.count() แยก SYSTEM/MOBILE, ทำ pagination และใช้ข้อความ “แสดงล่าสุด X จากทั้งหมด Y”'),
    ('P1', '10 URL ล่มแต่เคสค้าง 0 ดูขัดแย้ง', 'เคสถูกปิดได้เมื่อ backup ใช้งานได้ แต่การ์ดยังใช้คำว่า “ใช้ไม่ได้” สำหรับ URL หลักโดยไม่บอกเส้นทางสำรอง', 'เพิ่มสถานะ “ลิงก์หลักถูกบล็อก — ใช้งานผ่านสำรอง” และแสดง primary/backup คู่กัน'),
    ('P1', 'มิติยอดซิมไม่ชัด', 'Master Data แสดง 21 รายการตามบริษัท/ห้อง แต่หน้า agent แสดง 10 URL ไม่ซ้ำ ผู้ใช้จึงคิดว่ายอดผิด', 'แสดง 3 ตัวเลขพร้อมชื่อหน่วย: URL จริง, รายการ Master Data, ห้องที่ได้รับผลกระทบ'),
    ('P1', 'ข้อความยืนยัน 105/2 รอบทำให้เข้าใจผิด', 'failureStreak เป็นจำนวนเสียต่อเนื่องและเพิ่มเรื่อย ๆ แต่ UI ต่อท้าย /2 เหมือนยังรอยืนยัน', 'แสดง “ยืนยันแล้ว · เสียต่อเนื่อง 105 รอบ” และใช้ 1/2 เฉพาะก่อน confirm'),
    ('P2', 'รายงานรอบวัน 129 vs หน้า agent 136', 'report.ts กรองเฉพาะ URL หลักใน Master Data; agent jobs รวม URL สำรองเพิ่มอีก 7', 'ระบุ “URL หลัก 129” และ “URL สำรอง 7” หรือรวมพร้อมตัวแยกประเภท'),
    ('P2', 'React hydration errors บน Production', 'พบ minified React #418/#423/#425 ระหว่างเปลี่ยนหน้า; มีการ render Date.now() ใน client ทำให้ HTML server/client มีโอกาสไม่ตรง', 'คำนวณเวลา relative หลัง mount หรือส่ง now จาก server; เพิ่ม Playwright test ตรวจ console.error'),
    ('P2', 'ปุ่มคู่มือลอยทับ modal บนมือถือ', 'ตำแหน่ง fixed bottom-right ซ้อนกับเนื้อหา/ปุ่มปิดใน modal', 'ซ่อน FAB เมื่อ modal เปิด หรือย้ายเป็น header action พร้อม safe-area'),
    ('P3', 'Accessibility ของ modal/menu', 'modal รายละเอียดไม่มี role=dialog ที่ตรวจพบ และ overlay/X ใช้ชื่อ “ปิดเมนู” ซ้ำ', 'เพิ่ม role="dialog", aria-modal, focus trap และชื่อ accessible แยกกัน'),
]
for prio, title, evidence, fix in findings:
    t = doc.add_table(rows=1, cols=2)
    set_row_cant_split(t.rows[0])
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.rows[0].cells[0].width = Cm(1.5)
    t.rows[0].cells[1].width = Cm(16.0)
    fill = RED_LIGHT if prio == 'P1' else AMBER_LIGHT if prio == 'P2' else SLATE_LIGHT
    color = RED if prio == 'P1' else AMBER if prio == 'P2' else MUTED
    set_cell_fill(t.rows[0].cells[0], fill)
    p = t.rows[0].cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(prio); r.bold = True; r.font.color.rgb = RGBColor.from_string(color)
    set_cell_margins(t.rows[0].cells[0], top=100, bottom=100)
    set_cell_margins(t.rows[0].cells[1], top=100, bottom=100)
    p = t.rows[0].cells[1].paragraphs[0]
    r = p.add_run(title + '\n'); r.bold = True; r.font.size = Pt(10)
    r = p.add_run('หลักฐาน: ' + evidence + '\n'); r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(MUTED)
    r = p.add_run('ข้อเสนอ: ' + fix); r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

doc.add_page_break()

add_heading(doc, '5. ภาพหลักฐานจากระบบจริง', 1)
add_image(doc, '04-dashboard-counts.png', 'ภาพ 1 — Dashboard เวลา 09:11 น.: 651 = 496 เว็บไซต์ + 155 LINE และ 481 + 15 + 0 = 496')
add_image(doc, '01-agents-overview.png', 'ภาพ 2 — เครื่องตรวจ TRUE: ผลรวม URL ไม่ซ้ำ 136 รายการ (97/28/10/1)')

doc.add_page_break()
add_heading(doc, '5. ภาพหลักฐานจากระบบจริง (ต่อ)', 1)
add_image(doc, '02-result-detail.png', 'ภาพ 3 — Drill-down URL ใช้ไม่ได้: แสดงบริษัท ห้อง สาเหตุ SSL และเคสที่เกี่ยวข้อง')
add_image(doc, '03-incidents-history.png', 'ภาพ 4 — รายละเอียดเคสซิมที่ปิดแล้วผ่านลิงก์สำรอง พร้อมข้อมูลเครื่องและห้อง')

doc.add_page_break()
add_heading(doc, '5. ภาพหลักฐานจากระบบจริง (ต่อ)', 1)
add_image(doc, '12-incidents-counts.png', 'ภาพ 5 — หน้าเหตุการณ์แสดง “ประวัติทั้งหมด 236” ซึ่งเป็นจำนวนรายการที่โหลด ไม่ใช่ DB total')
add_image(doc, '11-kpi-counts.png', 'ภาพ 6 — KPI ระบบกลางนับ Incident จริง 764 เคส และปิดแล้ว 764 เคส')

doc.add_page_break()
add_heading(doc, '5. ภาพหลักฐานจากระบบจริง (ต่อ)', 1)
add_image(doc, '06-daily-report.png', 'ภาพ 7 — รายงานรอบวัน 29 ส.ค. 2569: 496 ลิงก์หลัก, 495 ใช้ได้, 1 ช้า, 0 ใช้ไม่ได้')
add_image(doc, '09-mobile-agents.png', 'ภาพ 8 — มุมมองมือถือหน้าเครื่องตรวจ: เมนู, ดาวน์โหลด APK, สร้าง QR และหลักการแจ้งเตือนอ่านได้', width=4.60)

add_heading(doc, '6. ผลตรวจทางเทคนิค', 1)
tech = doc.add_table(rows=1, cols=3)
tech.style = 'Table Grid'
for i, h in enumerate(['รายการ', 'ผล', 'รายละเอียด']):
    set_cell_fill(tech.rows[0].cells[i], BLUE)
    p = tech.rows[0].cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
technical_rows = [
    ('Automated tests', 'PASS 34/34', '0 failed; 412 ms'),
    ('TypeScript', 'PASS', 'npx tsc --noEmit --incremental false'),
    ('ESLint', 'PASS', 'No warnings or errors'),
    ('Next production build', 'PASS', 'Compiled, type-checked, generated 17 static pages/routes successfully'),
    ('Git consistency', 'PASS', 'main ตรง origin/main ที่ commit 28fa9d1; ไม่มี tracked change ระหว่าง QA'),
]
for item, status, detail in technical_rows:
    cells = tech.add_row().cells
    for i, val in enumerate([item, status, detail]):
        set_cell_margins(cells[i], top=90, bottom=90)
        cells[i].paragraphs[0].add_run(val)
    set_cell_fill(cells[1], GREEN_LIGHT)
    cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(GREEN)
    cells[1].paragraphs[0].runs[0].bold = True

add_heading(doc, '34 automated tests ครอบคลุมอะไร', 2)
for text in [
    'Runner: บันทึกผลครั้งแรก, ไม่เขียน log ซ้ำเมื่อสถานะเดิม, จำกัด concurrency และบันทึก transition',
    'State machine: ล่ม 2 รอบ, กู้คืน 2 รอบ, timeout เป็น inconclusive/slow, slow recovery, WAF/HTTP 503',
    'Mobile agent: enrollment, URL hash, redirect classification, backup recovery, ป้องกัน false outage จาก timeout',
    'สิทธิ์: ADMIN, ADMIN_LEAD, ADMIN_COMPANY, IT, MANAGEMENT และ SITE_STAFF',
    'Telegram: ข้อความล่ม/กู้คืน/กู้ผ่านสำรอง/สรุปรายวัน ระบุบริษัท ห้อง เครื่องและเคส',
    'รายงาน: deduplicate URL เดียวกันภายในบริษัทแม้มีหลายห้อง LINE',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)

add_heading(doc, 'ข้อจำกัดของการทดสอบ Production', 2)
for text in [
    'ไม่ได้กดส่ง Telegram จริง เพราะเป็นการส่งข้อความไปยังกลุ่มภายนอกและอาจรบกวนการปฏิบัติงาน; ทดสอบตัวสร้างข้อความและ route ด้วย automated tests แทน',
    'ไม่ได้สร้าง/ลบผู้ใช้ เครื่องตรวจ บริษัท หรือแก้ Master Data จริง; เปิดดู flow และตรวจ server permission จาก code/tests',
    'ไม่ได้กด “เช็คตอนนี้” ระหว่างรอบ QA เพื่อไม่สร้าง alert ใหม่โดยไม่จำเป็น; ตรวจ core checker ด้วย tests และอ่านค่าที่ cron/agent ส่งเข้ามาจริง',
    'ปุ่มบันทึก PNG และดาวน์โหลด APK แสดงใน UI แต่ browser automation ไม่ได้รับ download event จึงควรทำ manual smoke test บน Chrome/Android อีก 1 รอบ',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)

doc.add_page_break()
add_heading(doc, '7. แผนปรับปรุงที่แนะนำ', 1)
roadmap = [
    ('ทันที', 'แก้ incident count/pagination และเปลี่ยนป้าย 236 เป็น “แสดงล่าสุด … จากทั้งหมด …”'),
    ('ทันที', 'เปลี่ยนสถานะซิมให้แยก “URL หลักล่ม” ออกจาก “บริการใช้งานได้ผ่านสำรอง”'),
    ('ทันที', 'เปลี่ยน 105/2 เป็น “ยืนยันแล้ว · เสียต่อเนื่อง 105 รอบ”'),
    ('รอบถัดไป', 'แสดงยอดซิม 3 มิติ: unique URL / Master Data records / affected rooms'),
    ('รอบถัดไป', 'แยก URL หลัก 129 และ backup 7 ในรายงานรอบวัน'),
    ('รอบถัดไป', 'แก้ hydration errors และเพิ่ม E2E test ตรวจ console'),
    ('เก็บงาน UI', 'ซ่อนปุ่มคู่มือเมื่อเปิด modal และปรับ accessibility dialog/menu'),
    ('ก่อนส่งมอบ', 'manual smoke: APK install, QR enroll, PNG export, Print/PDF และ Telegram test ไปห้องทดสอบ'),
]
rt = doc.add_table(rows=1, cols=3)
rt.style = 'Table Grid'
for i, h in enumerate(['ลำดับ', 'ระยะ', 'งาน']):
    set_cell_fill(rt.rows[0].cells[i], BLUE)
    p = rt.rows[0].cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(h);r.bold=True;r.font.color.rgb=RGBColor(255,255,255)
for idx, (phase, item) in enumerate(roadmap, 1):
    cells=rt.add_row().cells
    vals=[str(idx),phase,item]
    for i,val in enumerate(vals):
        set_cell_margins(cells[i],top=85,bottom=85)
        p=cells[i].paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i<2 else WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(val)
    if phase=='ทันที': set_cell_fill(cells[1],RED_LIGHT)
    elif phase=='รอบถัดไป': set_cell_fill(cells[1],AMBER_LIGHT)
    else: set_cell_fill(cells[1],BLUE_LIGHT)

add_heading(doc, 'บทสรุป', 2)
closing = doc.add_paragraph()
shade_paragraph(closing, BLUE_LIGHT)
closing.paragraph_format.left_indent = Cm(0.25)
closing.paragraph_format.right_indent = Cm(0.25)
cr = closing.add_run('ระบบตรวจและแจ้งเตือนมีโครงสร้างหลักที่ดีและผ่าน test suite ทั้งหมด แต่การนำยอดไปใช้บริหารต้องแก้ “หน่วยนับ” และ “ยอดรวมจริง” ให้ชัดก่อน โดยเฉพาะหน้าเหตุการณ์และสถานะซิมที่มีลิงก์สำรอง เมื่อแก้ P1 ทั้ง 4 จุดแล้วควรรัน Regression QA และเก็บภาพ production ชุดใหม่อีกครั้ง')
cr.bold = True
cr.font.color.rgb = RGBColor.from_string(BLUE)

# Ensure all runs use a Thai-capable font, including runs created inside tables.
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        set_repeat_header_font(run)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_repeat_header_font(run)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
